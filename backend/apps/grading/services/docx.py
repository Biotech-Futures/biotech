"""DOCX renderers for marks summaries and participation certificates.

Templates live in two tiers:

  1. ``GradingSettings.marks_summary_template`` / ``certificate_template`` —
     client-provided, uploaded via the settings API. Takes precedence when
     present.
  2. Bundled fallbacks under ``apps/grading/templates/docx/*.docx`` — the
     client's real 2025 templates, so the endpoints produce the real
     documents out of the box.

Three template dialects are auto-detected per file:

  * ``<<[FieldName]>>`` text tokens — the client's marks release template
    (BTF 2025). Replaced run-aware so formatting and line breaks survive.
  * Word content controls with an alias (``firstName``/``lastName``/
    ``projectTitle``) — the client's merit certificate template.
  * Jinja ``{{ variable }}`` markers via ``docxtpl`` — anything else.
"""
from __future__ import annotations

import io
import re
import zipfile
from datetime import date
from decimal import Decimal, InvalidOperation
from pathlib import Path

from django.core.files.storage import default_storage
from docx import Document
from docx.oxml.ns import qn
from docxtpl import DocxTemplate

from ..models import GradingSettings


FALLBACK_DIR = Path(__file__).resolve().parent.parent / "templates" / "docx"

_TOKEN_RE = re.compile(r"<<\[(\w+)\]>>")


def _open_template(setting_field, fallback_filename: str):
    """Prefer the admin-uploaded docx; fall back to the in-repo template."""
    if setting_field:
        return default_storage.open(setting_field.name, "rb")
    fallback = FALLBACK_DIR / fallback_filename
    if not fallback.exists():
        raise FileNotFoundError(f"no template configured and fallback {fallback} missing")
    return fallback.open("rb")


def _document_xml(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        return z.read("word/document.xml").decode("utf8", errors="ignore")


def _has_angle_tokens(xml: str) -> bool:
    # Angle brackets inside text nodes are entity-escaped in the raw XML.
    return "&lt;&lt;[" in xml or "<<[" in xml


# ---------------------------------------------------------------------------
# <<[Field]>> token replacement


def _iter_paragraphs(container):
    """Yield every paragraph in the document body, descending into tables."""
    for p in container.paragraphs:
        yield p
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)


def _replace_tokens_in_paragraph(paragraph, fields: dict) -> None:
    """Replace ``<<[Name]>>`` tokens even when Word split them across runs.

    Works on the ``w:t`` text nodes directly: line breaks (``w:br``) and run
    formatting outside the token span are untouched. Unknown field names are
    blanked rather than left as visible template markers.
    """
    ts = paragraph._element.findall(f".//{qn('w:t')}")
    if not ts:
        return
    texts = [t.text or "" for t in ts]
    combined = "".join(texts)
    matches = list(_TOKEN_RE.finditer(combined))
    if not matches:
        return

    bounds = []
    pos = 0
    for txt in texts:
        bounds.append((pos, pos + len(txt)))
        pos += len(txt)

    # Right-to-left so earlier match offsets stay valid after each splice.
    for m in reversed(matches):
        value = str(fields.get(m.group(1), ""))
        s, e = m.span()
        start_i = end_i = None
        start_off = end_off = 0
        for i, (b0, b1) in enumerate(bounds):
            if b0 <= s < b1:
                start_i, start_off = i, s - b0
            if b0 < e <= b1:
                end_i, end_off = i, e - b0
        if start_i is None or end_i is None:
            continue
        if start_i == end_i:
            t = texts[start_i]
            texts[start_i] = t[:start_off] + value + t[end_off:]
        else:
            texts[start_i] = texts[start_i][:start_off] + value
            for j in range(start_i + 1, end_i):
                texts[j] = ""
            texts[end_i] = texts[end_i][end_off:]

    for t_el, new in zip(ts, texts):
        t_el.text = new


def _render_token_template(data: bytes, fields: dict) -> bytes:
    doc = Document(io.BytesIO(data))
    for paragraph in _iter_paragraphs(doc):
        _replace_tokens_in_paragraph(paragraph, fields)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Content-control (w:sdt alias) filling


def _fill_content_controls(data: bytes, fields: dict) -> bytes:
    doc = Document(io.BytesIO(data))
    for sdt in doc.element.body.iter(qn("w:sdt")):
        pr = sdt.find(qn("w:sdtPr"))
        alias = pr.find(qn("w:alias")) if pr is not None else None
        if alias is None:
            continue
        name = alias.get(qn("w:val"))
        if name not in fields:
            continue
        content = sdt.find(qn("w:sdtContent"))
        if content is None:
            continue
        ts = content.findall(f".//{qn('w:t')}")
        if not ts:
            continue
        ts[0].text = str(fields[name])
        for extra in ts[1:]:
            extra.text = ""
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Field mappings for the client's 2025 templates


def _sum_marks(criteria: list[dict]) -> Decimal:
    total = Decimal("0")
    for c in criteria:
        try:
            total += Decimal(c.get("mark") or "0")
        except InvalidOperation:
            continue
    return total


def _two_dp(value: Decimal) -> str:
    return str(value.quantize(Decimal("0.01")))


def marks_release_fields(context: dict) -> dict:
    """Flatten our context into the BTF marks release template's fields.

    Poster criteria map to P1..P10 in rubric order, SAQ to S1..S4. Fields we
    don't model (project title/category) render blank.
    """
    by_code = {c.get("code"): c for c in context.get("components", [])}

    def criteria(code: str) -> list[dict]:
        return (by_code.get(code) or {}).get("criteria", [])

    poster = criteria("POSTER")
    saq = criteria("SAQ")

    fields = {
        "TeamCode": context.get("group_name", ""),
        "ProjectTitle": context.get("project_title", ""),
        "ProjectCategoryHeading": "Project Category",
        "ProjectCategory": context.get("project_category", ""),
        "SolutionCategory": context.get("solution_category", ""),
        "Students": context.get("students", ""),
        "Mentor": context.get("mentors", ""),
        "SupervisorHeading": "Supervisor(s)",
        "Supervisors": context.get("supervisors", ""),
        "SchoolHeading": "School(s)",
        "Schools": context.get("schools", ""),
        "PosterComment": (by_code.get("POSTER") or {}).get("overall_comment", "")
        or context.get("poster_comment", ""),
    }
    for i in range(10):
        c = poster[i] if i < len(poster) else None
        fields[f"P{i + 1}"] = (c.get("mark") or "") if c else ""
        fields[f"P{i + 1}Comment"] = (c.get("comment") or "") if c else ""
    for i in range(4):
        c = saq[i] if i < len(saq) else None
        fields[f"S{i + 1}"] = (c.get("mark") or "") if c else ""
        fields[f"S{i + 1}Comment"] = (c.get("comment") or "") if c else ""

    poster_total = _sum_marks(poster)
    saq_total = _sum_marks(saq)
    fields["PosterTotal"] = _two_dp(poster_total)
    fields["SAQTotal"] = _two_dp(saq_total)
    fields["CombinedTotal"] = _two_dp(poster_total + saq_total)
    return fields


def certificate_fields(context: dict) -> dict:
    """Map our context onto the merit certificate's content-control aliases."""
    return {
        "firstName": context.get("first_name", ""),
        "lastName": context.get("last_name", ""),
        # No project-title field in the data model yet; the group name is the
        # closest identity we hold for the team's project.
        "projectTitle": context.get("project_title") or context.get("group_name", ""),
    }


# ---------------------------------------------------------------------------
# Public renderers


def render_marks_summary(context: dict) -> bytes:
    """Materialise a marks summary docx (see ``marks_summary_context``)."""
    settings = GradingSettings.load()
    with _open_template(settings.marks_summary_template, "marks_release.docx") as fh:
        data = fh.read()
    xml = _document_xml(data)
    if _has_angle_tokens(xml):
        return _render_token_template(data, marks_release_fields(context))
    doc = DocxTemplate(io.BytesIO(data))
    doc.render(context)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_participation_certificate(context: dict) -> bytes:
    """Materialise a certificate docx (see ``certificate_context``)."""
    settings = GradingSettings.load()
    with _open_template(settings.certificate_template, "merit_certificate.docx") as fh:
        data = fh.read()
    xml = _document_xml(data)
    fields = certificate_fields(context)
    if _has_angle_tokens(xml):
        return _render_token_template(data, fields)
    if "<w:sdt>" in xml or "w:alias" in xml:
        return _fill_content_controls(data, fields)
    doc = DocxTemplate(io.BytesIO(data))
    doc.render(context)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Context builders


def _team_details(group) -> dict:
    """Names of the group's active members, grouped by role, plus schools."""
    from apps.groups.models.group_members import GroupMembership

    memberships = (
        GroupMembership.objects.filter(group=group, left_at__isnull=True)
        .select_related("user")
        .order_by("joined_at")
    )
    students, mentors, supervisors, schools = [], [], [], []
    student_users = []
    for m in memberships:
        name = (m.user.get_full_name() or m.user.email).strip()
        role = m.membership_role
        if role == GroupMembership.MembershipRoleChoices.STUDENT:
            students.append(name)
            student_users.append(m.user_id)
        elif role == GroupMembership.MembershipRoleChoices.MENTOR:
            mentors.append(name)
        elif role == GroupMembership.MembershipRoleChoices.SUPERVISOR:
            supervisors.append(name)

    if student_users:
        from apps.users.models import StudentProfile

        for school in (
            StudentProfile.objects.filter(user_id__in=student_users)
            .exclude(school_name="")
            .exclude(school_name__isnull=True)
            .values_list("school_name", flat=True)
        ):
            if school not in schools:
                schools.append(school)

    return {
        "students": ", ".join(students),
        "mentors": ", ".join(mentors),
        "supervisors": ", ".join(supervisors),
        "schools": ", ".join(schools),
    }


def marks_summary_context(group, year: int, components: list[dict]) -> dict:
    settings = GradingSettings.load()
    return {
        "group_name": group.group_name,
        "year": year,
        "components": components,
        "director_1_name": settings.director_1_name or "",
        "director_2_name": settings.director_2_name or "",
        "generated_at": date.today().isoformat(),
        **_team_details(group),
    }


def certificate_context(
    student_full_name: str,
    group_name: str,
    year: int,
    first_name: str = "",
    last_name: str = "",
) -> dict:
    settings = GradingSettings.load()
    if not first_name and student_full_name:
        parts = student_full_name.split()
        first_name = parts[0]
        last_name = last_name or " ".join(parts[1:])
    return {
        "student_full_name": student_full_name,
        "first_name": first_name,
        "last_name": last_name,
        "group_name": group_name,
        "year": year,
        "director_1_name": settings.director_1_name or "",
        "director_2_name": settings.director_2_name or "",
        "issued_on": date.today().isoformat(),
    }
